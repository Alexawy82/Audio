"""
Document processor module for handling document text extraction and processing.
"""

import os
import logging
import re
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
import docx
import ebooklib
from ebooklib import epub
import PyPDF2
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing and text extraction."""
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger or logging.getLogger(__name__)
        
    def process_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple containing:
                - Extracted text content
                - Dictionary with metadata about the document
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
                
            extension = file_path.suffix.lower()
            
            if extension == '.docx':
                return self._process_docx(file_path)
            elif extension == '.pdf':
                return self._process_pdf(file_path)
            elif extension == '.epub':
                return self._process_epub(file_path)
            elif extension == '.txt':
                return self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            self.logger.error(f"Error processing document: {e}")
            raise
            
    def _process_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process a DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Process paragraphs with better formatting preservation
            processed_paragraphs = []
            
            # Extract content from all paragraphs including tables
            for para in doc.paragraphs:
                # Skip empty paragraphs
                if not para.text.strip():
                    continue
                    
                # Check if this is a heading
                if para.style.name.startswith('Heading'):
                    # Add extra newlines before headings for better separation
                    processed_paragraphs.append('\n' + para.text + '\n')
                else:
                    processed_paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        processed_paragraphs.append(' | '.join(row_text))
            
            # Combine paragraphs with proper newline separation  
            text = '\n\n'.join(processed_paragraphs)
            
            # Ensure document starts with the title if available
            doc_title = doc.core_properties.title or file_path.stem
            if doc_title and not text.startswith(doc_title):
                text = f"{doc_title}\n\n{text}"
                
            # Log text extraction stats
            self.logger.info(f"Extracted {len(text)} characters from DOCX with {len(doc.paragraphs)} paragraphs")
            
            metadata = {
                'title': doc_title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs)
            }
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
            
    def _process_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process a PDF file."""
        try:
            text = ''
            metadata = {
                'title': file_path.stem,
                'author': '',
                'created': None,
                'modified': None,
                'format': 'pdf',
                'page_count': 0
            }
            
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                metadata['page_count'] = len(pdf.pages)
                
                self.logger.info(f"Extracting text from PDF with {len(pdf.pages)} pages")
                
                # Extract text from each page
                total_text_extracted = 0
                for i, page in enumerate(pdf.pages):
                    self.logger.debug(f"Processing page {i+1}/{len(pdf.pages)}")
                    page_text = page.extract_text()
                    
                    # Skip completely empty pages
                    if not page_text or not page_text.strip():
                        self.logger.warning(f"Page {i+1} appears to be empty or contains only images")
                        continue
                    
                    total_text_extracted += len(page_text)
                        
                    # Add page number reference for better context
                    text += f"\n\n--- Page {i+1} ---\n\n"
                    text += page_text + '\n'
                    
                # Extract metadata if available
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('/Title', file_path.stem),
                        'author': pdf.metadata.get('/Author', ''),
                        'created': pdf.metadata.get('/CreationDate'),
                        'modified': pdf.metadata.get('/ModDate')
                    })
                
                # Verify that we extracted text successfully
                if not text.strip():
                    self.logger.error("No text extracted from PDF - the file may be scanned or contain only images")
                    text = f"This PDF document appears to contain no extractable text. It may be a scanned document or image-based PDF. The file has {len(pdf.pages)} pages. Please convert it to a text-based PDF for better results."
                else:
                    # Check if the extracted text was truncated
                    if total_text_extracted > len(text.strip()):
                        self.logger.warning(f"Text may have been truncated: extracted {total_text_extracted} chars but final text has {len(text.strip())} chars")
                    
                    self.logger.info(f"Successfully extracted {len(text)} characters from PDF")
                    
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing PDF: {e}")
            raise
            
    def _process_epub(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process an EPUB file."""
        try:
            book = epub.read_epub(file_path)
            full_text = []
            
            # Get book metadata
            title = book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else file_path.stem
            author = book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else ''
            
            # Start with book title and author
            full_text.append(title)
            if author:
                full_text.append(f"By {author}")
            full_text.append("\n")
            
            # Process items in document order
            item_count = 0
            extracted_content = set()  # Track extracted content to avoid duplication
            
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                item_count += 1
                self.logger.debug(f"Processing EPUB item {item_count}: {item.get_name()}")
                
                # Parse HTML content
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                
                # Process headings specifically to preserve structure
                for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    # Add extra newlines around headings for better chapter detection
                    heading_text = heading.get_text().strip()
                    if heading_text and heading_text not in extracted_content:
                        full_text.append(f"\n\n{heading_text}\n")
                        extracted_content.add(heading_text)
                
                # Process paragraphs and divs (common in EPUB)
                for element in soup.find_all(['p', 'div']):
                    # Skip if this is a container with other content we'll process separately
                    if element.find(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']) is not None:
                        continue
                        
                    element_text = element.get_text().strip()
                    if element_text and element_text not in extracted_content:
                        full_text.append(element_text)
                        extracted_content.add(element_text)
                
                # Process any remaining text by extracting the body text
                # This is a fallback to ensure we get all content
                body = soup.find('body')
                if body:
                    body_text = body.get_text().strip()
                    # Check if body text length is significantly more than what we've extracted
                    # This helps detect if we missed content due to unusual HTML structure
                    extracted_len = sum(len(text) for text in full_text[-20:] if text in body_text)
                    if len(body_text) > extracted_len * 1.5:  # 50% more content than what we extracted
                        self.logger.warning(f"EPUB item {item_count} had significant content missed in normal extraction")
                        # Split by newlines to preserve some structure
                        for line in body_text.split('\n'):
                            line = line.strip()
                            if line and line not in extracted_content:
                                full_text.append(line)
                                extracted_content.add(line)
                
                # Add section break
                full_text.append("\n")
            
            # Combine all text elements with appropriate spacing
            text = "\n\n".join(full_text)
            
            # Cleanup: Remove excessive newlines and whitespace
            text = re.sub(r'\n{3,}', '\n\n', text).strip()
            
            self.logger.info(f"Extracted {len(text)} characters from EPUB with {item_count} items")
            
            metadata = {
                'title': title,
                'author': author,
                'created': None,
                'modified': None,
                'format': 'epub',
                'item_count': item_count
            }
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing EPUB: {e}")
            raise
            
    def _process_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            metadata = {
                'title': file_path.stem,
                'author': '',
                'created': None,
                'modified': None,
                'format': 'txt'
            }
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing TXT: {e}")
            raise
            
    def split_into_chunks(self, text: str, max_chunk_size: int = 2000) -> List[str]:
        """Split text into chunks of specified size."""
        try:
            chunks = []
            current_chunk = ''
            
            # Split by paragraphs first
            paragraphs = text.split('\n\n')
            
            for paragraph in paragraphs:
                # If paragraph is too long, split by sentences
                if len(paragraph) > max_chunk_size:
                    sentences = paragraph.split('. ')
                    current_sentence = ''
                    
                    for sentence in sentences:
                        if len(current_sentence) + len(sentence) + 2 <= max_chunk_size:
                            current_sentence += sentence + '. '
                        else:
                            if current_sentence:
                                chunks.append(current_sentence.strip())
                            current_sentence = sentence + '. '
                            
                    if current_sentence:
                        chunks.append(current_sentence.strip())
                        
                # If current chunk + paragraph is too long, start new chunk
                elif len(current_chunk) + len(paragraph) + 2 > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = paragraph + '\n\n'
                    
                else:
                    current_chunk += paragraph + '\n\n'
                    
            if current_chunk:
                chunks.append(current_chunk.strip())
                
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error splitting text into chunks: {e}")
            raise 