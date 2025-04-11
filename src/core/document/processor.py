"""
Document processing module for DocToAudiobook.
"""

import os
import logging
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import docx
import PyPDF2
from ..models.document import Document, DocumentMetadata
from ..utils.error import ErrorHandler

class DocumentProcessor:
    """Handles document processing operations."""
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        """Initialize document processor."""
        self.logger = logger or logging.getLogger(__name__)
        self.error_handler = ErrorHandler()
        
    def process_document(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process a document file and extract text and metadata."""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
                
            # Get file extension
            ext = file_path.suffix.lower()
            
            # Process based on file type
            if ext == '.docx':
                return self._process_docx(file_path)
            elif ext == '.pdf':
                return self._process_pdf(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
                
        except Exception as e:
            self.error_handler.log_error(e, {'file_path': str(file_path)})
            raise
            
    def _process_docx(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Process a DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract metadata
            metadata = DocumentMetadata(
                title=doc.core_properties.title or file_path.stem,
                author=doc.core_properties.author or 'Unknown',
                created=doc.core_properties.created,
                modified=doc.core_properties.modified,
                pages=len(doc.paragraphs)
            )
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX file: {e}")
            raise
            
    def _process_pdf(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Process a PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                # Extract text
                text = ''
                for page in pdf.pages:
                    text += page.extract_text() + '\n'
                    
                # Extract metadata
                metadata = DocumentMetadata(
                    title=pdf.metadata.get('/Title', file_path.stem) if pdf.metadata else file_path.stem,
                    author=pdf.metadata.get('/Author', 'Unknown') if pdf.metadata else 'Unknown',
                    created=None,  # PDFs don't always store creation date
                    modified=None,  # PDFs don't always store modification date
                    pages=len(pdf.pages)
                )
                
                return text, metadata
                
        except Exception as e:
            self.logger.error(f"Error processing PDF file: {e}")
            raise 